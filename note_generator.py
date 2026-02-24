#!/usr/bin/env python3
"""
Note Generator: Given a lecture/speech transcription (txt, pdf, or json),
generates a comprehensive note document as a .docx file using an external LLM.
Output filename: input name + "_note" before extension (e.g., lecture.txt -> lecture_note.docx).
Takes input file path and API key as arguments.
"""

import argparse
import json
import os
import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Install python-docx: pip install python-docx", file=sys.stderr)
    sys.exit(1)

try:
    from openai import OpenAI
except ImportError:
    OpenAI = None

try:
    from pypdf import PdfReader
except ImportError:
    PdfReader = None


def read_txt(path: Path) -> str:
    """Read plain text file."""
    return path.read_text(encoding="utf-8", errors="replace").strip()


def read_pdf(path: Path) -> str:
    """Extract text from PDF."""
    if PdfReader is None:
        raise ImportError("PDF support requires pypdf: pip install pypdf")
    reader = PdfReader(str(path))
    parts = []
    for page in reader.pages:
        parts.append(page.extract_text() or "")
    return "\n".join(parts).strip()


def read_json(path: Path) -> str:
    """Extract text from JSON (tries common keys: text, transcript, content)."""
    raw = path.read_text(encoding="utf-8", errors="replace")
    data = json.loads(raw)
    if isinstance(data, str):
        return data.strip()
    if isinstance(data, list):
        # List of segments or paragraphs
        return "\n".join(
            (x if isinstance(x, str) else x.get("text", str(x)) for x in data)
        ).strip()
    for key in ("text", "transcript", "content", "body"):
        if key in data and data[key]:
            val = data[key]
            if isinstance(val, str):
                return val.strip()
            if isinstance(val, list):
                return "\n".join(
                    (x if isinstance(x, str) else str(x) for x in val)
                ).strip()
    # Fallback: join all string values
    return "\n".join(str(v) for v in data.values() if v and isinstance(v, str)).strip()


def load_transcription(path: Path) -> str:
    """Load transcription from .txt, .pdf, or .json."""
    path = Path(path)
    if not path.exists():
        raise FileNotFoundError(f"File not found: {path}")
    suf = path.suffix.lower()
    if suf == ".txt":
        return read_txt(path)
    if suf == ".pdf":
        return read_pdf(path)
    if suf == ".json":
        return read_json(path)
    raise ValueError(f"Unsupported format: {suf}. Use .txt, .pdf, or .json.")


NOTES_SYSTEM_PROMPT = """You are a helpful assistant that turns lecture or speech transcriptions into clear, comprehensive notes.

Output your response using exactly these section headers (copy them as-is):
## Summary
## Main Points
## Detailed Notes

- Under "Summary", write a concise overview in 1–3 paragraphs.
- Under "Main Points", list the key takeaways as bullet points (one per line, start each with "- ").
- Under "Detailed Notes", write well-organized notes that capture the main content, key definitions, examples, and structure of the lecture. Use paragraphs and optional subheadings. Do not simply copy the transcript; summarize and structure it for study.

Use only the section headers above. Do not add other sections or markdown beyond simple bullets and paragraphs."""


def generate_notes_with_llm(transcription: str, api_key: str, model: str = "gpt-4o-mini", base_url: str | None = None) -> dict:
    """Call external LLM to generate structured notes from transcription."""
    if OpenAI is None:
        raise ImportError("LLM support requires openai: pip install openai")
    client = OpenAI(api_key=api_key, base_url=base_url or os.environ.get("OPENAI_API_BASE"))
    # Truncate if very long to stay within context limits
    max_chars = 120_000
    if len(transcription) > max_chars:
        transcription = transcription[:max_chars] + "\n\n[... transcription truncated ...]"
    response = client.chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": NOTES_SYSTEM_PROMPT},
            {"role": "user", "content": f"Generate comprehensive notes from this lecture/speech transcription:\n\n{transcription}"},
        ],
        temperature=0.3,
    )
    raw = (response.choices[0].message.content or "").strip()
    return parse_llm_notes(raw)


def parse_llm_notes(raw: str) -> dict:
    """Parse LLM output into summary, main_points, and paragraphs for docx."""
    summary = ""
    main_points: list[str] = []
    paragraphs: list[str] = []
    section = None
    current_para: list[str] = []
    for line in raw.split("\n"):
        if re.match(r"^##\s*Summary\s*$", line, re.IGNORECASE):
            section = "summary"
            continue
        if re.match(r"^##\s*Main Points\s*$", line, re.IGNORECASE):
            section = "main_points"
            if current_para:
                paragraphs.append("\n".join(current_para).strip())
                current_para = []
            continue
        if re.match(r"^##\s*Detailed Notes\s*$", line, re.IGNORECASE):
            section = "paragraphs"
            if current_para:
                paragraphs.append("\n".join(current_para).strip())
                current_para = []
            continue
        if section == "summary":
            summary = (summary + "\n" + line).strip()
        elif section == "main_points":
            stripped = line.strip()
            if stripped.startswith("- "):
                main_points.append(stripped[2:].strip())
            elif stripped:
                main_points.append(stripped)
        elif section == "paragraphs":
            stripped = line.strip()
            if not stripped:
                if current_para:
                    paragraphs.append("\n".join(current_para).strip())
                    current_para = []
            else:
                current_para.append(stripped)
    if current_para:
        paragraphs.append("\n".join(current_para).strip())
    # If parsing found nothing, treat whole response as detailed notes
    if not summary and not main_points and not paragraphs and raw:
        paragraphs = [p.strip() for p in raw.split("\n\n") if p.strip()]
        if not paragraphs:
            paragraphs = [raw]
    return {"summary": summary, "main_points": main_points, "paragraphs": paragraphs}


def write_docx(notes: dict, output_path: Path, source_name: str) -> None:
    """Write structured notes to a Word document."""
    doc = Document()
    title = doc.add_heading(f"Notes: {source_name}", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("Summary", level=1)
    doc.add_paragraph(notes["summary"] or "(No summary generated.)")

    doc.add_heading("Main Points", level=1)
    for point in notes["main_points"]:
        doc.add_paragraph(point, style="List Bullet")

    doc.add_heading("Detailed Notes", level=1)
    for p in notes["paragraphs"]:
        if p not in notes["main_points"]:  # avoid duplicate
            doc.add_paragraph(p)

    doc.save(str(output_path))


def output_path_for(input_path: Path) -> Path:
    """Return output path: input name + '_note' before extension, .docx."""
    base = input_path.parent / input_path.stem
    return base.with_name(base.name + "_note.docx")


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Generate comprehensive notes (docx) from a lecture/speech transcription using an external LLM."
    )
    parser.add_argument(
        "input_file",
        type=Path,
        help="Path to transcription file (.txt, .pdf, or .json)",
    )
    parser.add_argument(
        "--api-key",
        "-k",
        type=str,
        required=True,
        help="AI (API) key for the LLM service",
    )
    parser.add_argument(
        "-o",
        "--output",
        type=Path,
        default=None,
        help="Output .docx path (default: input name + '_note.docx')",
    )
    parser.add_argument(
        "--model",
        type=str,
        default="gpt-4o-mini",
        help="LLM model name (default: gpt-4o-mini)",
    )
    args = parser.parse_args()

    input_path = args.input_file
    api_key = args.api_key.strip()
    if not api_key:
        print("Error: API key cannot be empty.", file=sys.stderr)
        sys.exit(1)

    try:
        text = load_transcription(input_path)
    except (FileNotFoundError, ValueError, ImportError) as e:
        print(f"Error: {e}", file=sys.stderr)
        sys.exit(1)

    if not text:
        print("Error: No text content found in the input file.", file=sys.stderr)
        sys.exit(1)

    print("Generating notes with LLM...", file=sys.stderr)
    try:
        notes = generate_notes_with_llm(text, api_key, model=args.model)
    except Exception as e:
        print(f"Error calling LLM: {e}", file=sys.stderr)
        sys.exit(1)

    out = args.output if args.output is not None else output_path_for(input_path)
    out = Path(out)
    if out.suffix.lower() != ".docx":
        out = out.with_suffix(".docx")

    write_docx(notes, out, input_path.stem)
    print(f"Notes saved to: {out.resolve()}")


if __name__ == "__main__":
    main()
